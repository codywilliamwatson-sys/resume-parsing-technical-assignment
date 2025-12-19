"""Tests for file parsers."""

import pytest
from pathlib import Path
from unittest.mock import Mock, patch, mock_open
import pypdf
from docx import Document

from src.parsers.pdf_parser import PDFParser
from src.parsers.word_parser import WordParser


class TestPDFParser:
    """Test cases for PDFParser."""
    
    def test_init(self):
        """Test PDFParser initialization."""
        parser = PDFParser()
        assert parser._supported_extensions == ['.pdf']
    
    def test_get_supported_extensions(self):
        """Test getting supported extensions."""
        parser = PDFParser()
        extensions = parser.get_supported_extensions()
        assert extensions == ['.pdf']
        # Ensure it returns a copy
        extensions.append('.docx')
        assert parser.get_supported_extensions() == ['.pdf']
    
    def test_can_parse_pdf_file(self):
        """Test can_parse returns True for PDF files."""
        parser = PDFParser()
        file_path = Path("test.pdf")
        assert parser.can_parse(file_path) is True
    
    def test_can_parse_non_pdf_file(self):
        """Test can_parse returns False for non-PDF files."""
        parser = PDFParser()
        file_path = Path("test.docx")
        assert parser.can_parse(file_path) is False
    
    def test_can_parse_case_insensitive(self):
        """Test can_parse is case insensitive."""
        parser = PDFParser()
        assert parser.can_parse(Path("test.PDF")) is True
        assert parser.can_parse(Path("test.Pdf")) is True
    
    def test_validate_file_not_found(self):
        """Test validate_file raises FileNotFoundError for non-existent file."""
        parser = PDFParser()
        file_path = Path("nonexistent.pdf")
        
        with pytest.raises(FileNotFoundError):
            parser.validate_file(file_path)
    
    def test_validate_file_is_directory(self, tmp_path):
        """Test validate_file raises ValueError for directory."""
        parser = PDFParser()
        dir_path = tmp_path / "test_dir"
        dir_path.mkdir()
        
        with pytest.raises(ValueError, match="Path is not a file"):
            parser.validate_file(dir_path)
    
    def test_validate_file_empty(self, tmp_path):
        """Test validate_file raises ValueError for empty file."""
        parser = PDFParser()
        file_path = tmp_path / "empty.pdf"
        file_path.touch()
        
        with pytest.raises(ValueError, match="File is empty"):
            parser.validate_file(file_path)
    
    @patch('builtins.open', new_callable=mock_open)
    @patch('src.parsers.pdf_parser.pypdf.PdfReader')
    def test_parse_success(self, mock_pdf_reader, mock_file, tmp_path):
        """Test successful PDF parsing."""
        parser = PDFParser()
        file_path = tmp_path / "test.pdf"
        file_path.write_bytes(b"fake pdf content")
        
        # Mock PDF reader
        mock_page1 = Mock()
        mock_page1.extract_text.return_value = "Page 1 content"
        mock_page2 = Mock()
        mock_page2.extract_text.return_value = "Page 2 content"
        
        mock_reader_instance = Mock()
        mock_reader_instance.pages = [mock_page1, mock_page2]
        mock_pdf_reader.return_value = mock_reader_instance
        
        result = parser.parse(file_path)
        
        assert result == "Page 1 content\nPage 2 content"
        mock_file.assert_called_once()
        mock_pdf_reader.assert_called_once()
    
    @patch('builtins.open', new_callable=mock_open)
    @patch('src.parsers.pdf_parser.pypdf.PdfReader')
    def test_parse_invalid_pdf(self, mock_pdf_reader, mock_file, tmp_path):
        """Test parsing invalid PDF raises ValueError."""
        parser = PDFParser()
        file_path = tmp_path / "invalid.pdf"
        file_path.write_bytes(b"invalid content")
        
        mock_pdf_reader.side_effect = pypdf.errors.PdfReadError("Invalid PDF")
        
        with pytest.raises(ValueError, match="Invalid PDF file"):
            parser.parse(file_path)
    
    @patch('builtins.open', new_callable=mock_open)
    def test_parse_io_error(self, mock_file, tmp_path):
        """Test parsing raises IOError on file read error."""
        parser = PDFParser()
        file_path = tmp_path / "test.pdf"
        file_path.write_bytes(b"content")
        
        mock_file.side_effect = IOError("Cannot read file")
        
        with pytest.raises(IOError, match="Error reading PDF file"):
            parser.parse(file_path)


class TestWordParser:
    """Test cases for WordParser."""
    
    def test_init(self):
        """Test WordParser initialization."""
        parser = WordParser()
        assert parser._supported_extensions == ['.docx']
    
    def test_get_supported_extensions(self):
        """Test getting supported extensions."""
        parser = WordParser()
        extensions = parser.get_supported_extensions()
        assert extensions == ['.docx']
        # Ensure it returns a copy
        extensions.append('.pdf')
        assert parser.get_supported_extensions() == ['.docx']
    
    def test_can_parse_docx_file(self):
        """Test can_parse returns True for DOCX files."""
        parser = WordParser()
        file_path = Path("test.docx")
        assert parser.can_parse(file_path) is True
    
    def test_can_parse_non_docx_file(self):
        """Test can_parse returns False for non-DOCX files."""
        parser = WordParser()
        file_path = Path("test.pdf")
        assert parser.can_parse(file_path) is False
    
    def test_can_parse_case_insensitive(self):
        """Test can_parse is case insensitive."""
        parser = WordParser()
        assert parser.can_parse(Path("test.DOCX")) is True
        assert parser.can_parse(Path("test.Docx")) is True
    
    def test_validate_file_not_found(self):
        """Test validate_file raises FileNotFoundError for non-existent file."""
        parser = WordParser()
        file_path = Path("nonexistent.docx")
        
        with pytest.raises(FileNotFoundError):
            parser.validate_file(file_path)
    
    def test_validate_file_is_directory(self, tmp_path):
        """Test validate_file raises ValueError for directory."""
        parser = WordParser()
        dir_path = tmp_path / "test_dir"
        dir_path.mkdir()
        
        with pytest.raises(ValueError, match="Path is not a file"):
            parser.validate_file(dir_path)
    
    def test_validate_file_empty(self, tmp_path):
        """Test validate_file raises ValueError for empty file."""
        parser = WordParser()
        file_path = tmp_path / "empty.docx"
        file_path.touch()
        
        with pytest.raises(ValueError, match="File is empty"):
            parser.validate_file(file_path)
    
    @patch('src.parsers.word_parser.Document')
    def test_parse_success_with_paragraphs(self, mock_document, tmp_path):
        """Test successful Word document parsing with paragraphs."""
        parser = WordParser()
        file_path = tmp_path / "test.docx"
        file_path.write_bytes(b"fake docx content")
        
        # Mock document
        mock_doc = Mock()
        mock_doc.paragraphs = [
            Mock(text="Paragraph 1"),
            Mock(text=""),
            Mock(text="Paragraph 2"),
            Mock(text="   "),  # Whitespace only
        ]
        mock_doc.tables = []
        mock_document.return_value = mock_doc
        
        result = parser.parse(file_path)
        
        assert result == "Paragraph 1\nParagraph 2"
        mock_document.assert_called_once_with(file_path)
    
    @patch('src.parsers.word_parser.Document')
    def test_parse_success_with_tables(self, mock_document, tmp_path):
        """Test successful Word document parsing with tables."""
        parser = WordParser()
        file_path = tmp_path / "test.docx"
        file_path.write_bytes(b"fake docx content")
        
        # Mock document with tables
        mock_doc = Mock()
        mock_doc.paragraphs = [Mock(text="Paragraph 1")]
        
        # Mock table
        mock_cell1 = Mock(text="Cell 1")
        mock_cell2 = Mock(text="  Cell 2  ")
        mock_cell3 = Mock(text="")
        mock_row = Mock(cells=[mock_cell1, mock_cell2, mock_cell3])
        mock_table = Mock(rows=[mock_row])
        mock_doc.tables = [mock_table]
        
        mock_document.return_value = mock_doc
        
        result = parser.parse(file_path)
        
        assert "Paragraph 1" in result
        assert "Cell 1 | Cell 2" in result
        mock_document.assert_called_once_with(file_path)
    
    @patch('src.parsers.word_parser.Document')
    def test_parse_success_with_paragraphs_and_tables(self, mock_document, tmp_path):
        """Test successful Word document parsing with both paragraphs and tables."""
        parser = WordParser()
        file_path = tmp_path / "test.docx"
        file_path.write_bytes(b"fake docx content")
        
        mock_doc = Mock()
        mock_doc.paragraphs = [Mock(text="Paragraph 1")]
        
        mock_row = Mock(cells=[Mock(text="Cell 1"), Mock(text="Cell 2")])
        mock_table = Mock(rows=[mock_row])
        mock_doc.tables = [mock_table]
        
        mock_document.return_value = mock_doc
        
        result = parser.parse(file_path)
        
        assert result == "Paragraph 1\n\nCell 1 | Cell 2"
    
    @patch('src.parsers.word_parser.Document')
    def test_parse_io_error(self, mock_document, tmp_path):
        """Test parsing raises IOError on document read error."""
        parser = WordParser()
        file_path = tmp_path / "test.docx"
        file_path.write_bytes(b"content")
        
        mock_document.side_effect = IOError("Cannot read document")
        
        with pytest.raises(IOError, match="Error reading Word document"):
            parser.parse(file_path)
    
    @patch('src.parsers.word_parser.Document')
    def test_parse_preserves_file_not_found_error(self, mock_document, tmp_path):
        """Test parsing preserves FileNotFoundError."""
        parser = WordParser()
        file_path = tmp_path / "test.docx"
        file_path.write_bytes(b"content")
        
        mock_document.side_effect = FileNotFoundError("File not found")
        
        with pytest.raises(FileNotFoundError):
            parser.parse(file_path)
    
    @patch('src.parsers.word_parser.Document')
    def test_parse_preserves_value_error(self, mock_document, tmp_path):
        """Test parsing preserves ValueError."""
        parser = WordParser()
        file_path = tmp_path / "test.docx"
        file_path.write_bytes(b"content")
        
        mock_document.side_effect = ValueError("Invalid value")
        
        with pytest.raises(ValueError):
            parser.parse(file_path)

