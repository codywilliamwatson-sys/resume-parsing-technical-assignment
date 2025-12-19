"""
Word document parser implementation.
Handles parsing of Microsoft Word (.docx) resume files.
"""

import logging
from pathlib import Path
from docx import Document

from .file_parser import FileParser

logger = logging.getLogger(__name__)


class WordParser(FileParser):
    """Parser for Microsoft Word (.docx) files."""
    
    def __init__(self):
        """Initialize the Word parser."""
        self._supported_extensions = ['.docx']
    
    def parse(self, file_path: Path) -> str:
        """
        Parse a Word document and extract text content.
        
        Args:
            file_path: Path to the Word document to be parsed
            
        Returns:
            String containing the extracted text content from paragraphs and tables
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If the file format is invalid or unsupported
            IOError: If there's an error reading the file
        """
        logger.info(f"Parsing Word document: {file_path}")
        self.validate_file(file_path)
        
        try:
            doc = Document(file_path)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(' | '.join(row_text))
            
            all_text = '\n'.join(text_content)
            if table_text:
                all_text += '\n\n' + '\n'.join(table_text)
            
            logger.info(f"Successfully parsed Word document: {len(text_content)} paragraphs, {len(table_text)} table rows, {len(all_text)} characters extracted")
            return all_text
        except Exception as e:
            if isinstance(e, (FileNotFoundError, ValueError)):
                raise
            logger.error(f"Error reading Word document: {file_path} - {str(e)}")
            raise IOError(f"Error reading Word document: {file_path}") from e
    
    def can_parse(self, file_path: Path) -> bool:
        """
        Check if this parser can handle the given file.
        
        Args:
            file_path: Path to the file to check
            
        Returns:
            True if the file has a .docx extension, False otherwise
        """
        return file_path.suffix.lower() in self._supported_extensions
    
    def get_supported_extensions(self) -> list[str]:
        """
        Get the list of file extensions this parser supports.
        
        Returns:
            List containing ['.docx']
        """
        return self._supported_extensions.copy()
    

